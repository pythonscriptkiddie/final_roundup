#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu May 23 13:45:44 2019

@author: thomassullivan
"""
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from objects import Article

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    #print(paragraph)
    #print(text)
    #print(url)
    try:
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    
        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run ()
        r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True
        
        return hyperlink
    except Exception as e:
        print(e)


def add_article(document, article):
    #print(article)
    try:
        new_paragraph = document.add_paragraph('') #add blank paragraph that we append the text to
        add_hyperlink(paragraph=new_paragraph, text=article.name, url=article.link)
        #print(Article.get_date_formatted(article))
        new_paragraph.add_run(' ({0}) '.format(Article.get_date_formatted(article))) #blank space between the link and the description
        new_paragraph.add_run(article.description)
    except Exception as e:
        print(e)

def add_section(document, section):
    section_name = document.add_paragraph(section.section_name)
    section.categories.sort(key=lambda x: x.name, reverse=True)
    section.categories.reverse()
    for category in section.categories:
        add_category(document, category)
    
def add_category(document, category):
    category_name = document.add_paragraph(category.category_name)
    #category.articles = category.articles.sort()
    category.articles.sort(key=lambda x: x.name, reverse=True)
    category.articles.reverse()
    for article in category.articles:
        #print(article)
        add_article(document, article)
  

def create_roundup2(document, roundup_title, categories):
    title = document.add_paragraph(roundup_title)
    for category in categories:
        add_category(document, category)
        
def complete_roundup2(filename, roundup_title, sections):
    new_document = docx.Document()
    create_roundup2(new_document, roundup_title, sections)
    new_document.save('{0}.docx'.format(filename))
    
      
def create_roundup_docx(document, roundup_title, categories):
    title = document.add_paragraph(roundup_title)
    for category in categories:
        add_category(document, category)
        
def create_complete_roundup(filename, roundup_title, categories):
    new_document = docx.Document()
    create_roundup_docx(new_document, roundup_title, categories)
    new_document.save('{0}.docx'.format(filename))
    

if __name__ == '__main__':
    
    print('roundup_docx2 loaded')
    
